#-*- coding: UTF-8 -*- 
'''

@author: yanghp

'''
from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
import datetime

curDate= datetime.datetime.now().strftime("%Y/%m/%d")
document = Document()
#加入标题
titleName=u'数据仓库运行日报('+curDate+')'
document.add_heading(titleName,0).paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER


paragraph =document.add_heading(u'二级标题',1)
paragraph_format = paragraph.paragraph_format
paragraph_format.left_indent = Inches(0.3)

document.add_heading(u'二级标题',2)


#添加文本
paragraph = document.add_paragraph(u'添加了文本')
paragraph.paragraph_format.alignment = WD_TAB_ALIGNMENT.CENTER
#设置字号
run = paragraph.add_run(u'设置字号')
run.font.size=Pt(24)

#设置字体
run = paragraph.add_run('Set Font,')
run.font.name='Consolas'
#设置中文字体
run = paragraph.add_run(u'设置中文字体，')
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

#设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True

#设置粗体
run = paragraph.add_run(u'粗体').bold = True

#增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

#增加有序列表
document.add_paragraph(
    u'有序列表元素1',style='List Number'
)
document.add_paragraph(
    u'有序列别元素2',style='List Number'
)

#增加无序列表
document.add_paragraph(
    u'无序列表元素1',style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2',style='List Bullet'
)

#增加图片（此处使用相对位置）
#document.add_picture('jdb.jpg',width=Inches(1.25))

#增加表格
table = document.add_table(rows=3,cols=3)
hdr_cells=table.rows[0].cells
hdr_cells[0].text=u"第一列"
hdr_cells[1].text=u"第二列"
hdr_cells[2].text=u"第三列"

hdr_cells = table.rows[1].cells
hdr_cells[0].text =u'2'
hdr_cells[1].text = u'aerszvfdgx'
hdr_cells[2].text = u'abdzfgxfdf'

hdr_cells = table.rows[2].cells
hdr_cells[0].text = u'3'
hdr_cells[1].text = u'cafdwvaef'
hdr_cells[2].text = u'aabs zfgf'

#增加分页
document.add_page_break()
#保存文件
document.save('E:\demo.docx')
